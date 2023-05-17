import os
import shutil
from itertools import cycle
from collections import namedtuple
import random
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 段落居中
from docx.oxml.ns import qn
from docx.shared import Pt
from docx import Document
from openpyxl import load_workbook

import logging

logger = logging.getLogger("")
logger.setLevel(logging.DEBUG)
fh = logging.FileHandler("output.log")
fh.setLevel(logging.DEBUG)
ch = logging.StreamHandler()
ch.setLevel(logging.INFO)
formatter = logging.Formatter("%(asctime)s - %(name)s - %(levelname)s - %(message)s")
ch.setFormatter(formatter)
fh.setFormatter(formatter)
logger.addHandler(ch)
logger.addHandler(fh)


Row = namedtuple(
    "Row",
    [
        "sn",
        "type",
        "level",
        "question",
        "choice_a",
        "choice_b",
        "choice_c",
        "choice_d",
        "choice_e",
        "answer",
        "reference",
        "category",
    ],
)


# 如xlsx题库文件所在目录存在output文件夹则删除并新建
def init_output_folder(xlsx_path):
    logger.info(os.path.dirname(xlsx_path))
    output_path = os.path.join(os.path.dirname(xlsx_path), "output")
    if os.path.exists(output_path):
        shutil.rmtree(output_path)
    os.mkdir(output_path)
    return output_path


# 打乱选项
def mess_up_choices(row: Row):
    trans = str.maketrans("ABCDE", "01234")
    correct_list = list(map(int, list(row.answer.translate(trans))))
    choices = list(row[4:9])
    length = 5 if choices[-1] else 4
    temp = [choices[i] for i in range(length) if i not in correct_list]
    temp += [""] * (length - len(temp))
    random.shuffle(temp)
    random.shuffle(correct_list)
    for idx, i in enumerate(correct_list):
        j = temp.index("")
        temp[j] = choices[i]
        correct_list[idx] = j
    answer = "".join(map(str, correct_list)).translate({v: k for k, v in trans.items()})
    new_row = list(row)
    temp += [None] * (5 - len(temp)) + [
        answer,
    ]
    new_row[4:10] = temp
    return Row(*new_row)


def add_sheet_part(
    sheet, q_doc, a_doc, mumber_of_questions, is_choice=True, is_mess_up=True
):
    # 生成来源行号字典
    category_dict = dict()
    chosen_q_set = set()
    a_para = ""
    for row_no, cells in enumerate(sheet.iter_rows(2), 2):  # 跳过表头
        r = Row(*[c.value for c in cells])
        if r.category:  # 防止输入空行
            if r.category in category_dict:
                category_dict[r.category].append(row_no)
            else:
                category_dict.update(
                    {
                        r.category: [
                            row_no,
                        ]
                    }
                )
    logger.info(category_dict)
    # 根据来源循环抽题
    for cat in cycle(category_dict.keys()):
        random.shuffle(category_dict[cat])
        try:
            chosen_q_set.add(category_dict[cat].pop())
        except IndexError:  # 来源抽完跳过
            continue
        if len(chosen_q_set) == mumber_of_questions:
            break
    logger.info(chosen_q_set)
    for idx, i in enumerate(chosen_q_set, 1):
        r = Row(*[c.value for c in sheet[i]])
        if is_choice:
            if is_mess_up:
                r = mess_up_choices(r)
            formated_choice_e = "   E、{}\n".format(r.choice_e) if r.choice_e else ""
            q_para = "{0}、{1}\n   A、{2}\n   B、{3}\n   C、{4}\n   D、{5}\n{6}".format(
                idx,
                r.question,
                r.choice_a,
                r.choice_b,
                r.choice_c,
                r.choice_d,
                formated_choice_e,
            )
        else:
            q_para = "{0}、{1}".format(idx, r.question)
        q_doc.add_paragraph(q_para)

        a_para += "{0}、{1}   ".format(idx, r.answer)
    a_doc.add_paragraph(a_para)  # 答案


def create(xlsx_path, number_of_copies, is_mess_up):
    wb = load_workbook(xlsx_path, data_only=True)

    dan_sheet, duo_sheet, tian_sheet, pan_sheet, jian_sheet = wb.worksheets

    output_path = init_output_folder(xlsx_path)

    for count in range(number_of_copies):
        # 打开文档
        q_doc = Document()
        a_doc = Document()

        # 修改正文的中文字体类型，示例代码：（全局设置）
        q_doc.styles["Normal"].font.name = "仿宋"
        q_doc.styles["Normal"].font.size = Pt(16)  # 16对应仿宋三号
        q_doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

        head_str = "题库随机抽题考试" + str(count) + "\n"  # 文章标题
        # document.add_heading(head_str,0)
        head_paragraph = q_doc.add_paragraph("")  # 添加一个段落
        head_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落居中
        run = head_paragraph.add_run(head_str)
        run.bold = True  # 粗体是
        run.italic = False  # 斜体否

        q_doc.styles["Normal"].font.size = Pt(12)  # 12对应仿宋小四号

        # ------------------------------*************************------------------------------
        dan_para = "第一部分  单选题(8题)"
        q_doc.add_paragraph(dan_para)  # 考题
        a_doc.add_paragraph(dan_para)  # 答案

        add_sheet_part(dan_sheet, q_doc, a_doc, 8, is_mess_up=is_mess_up)

        # ------------------------------*************************------------------------------
        duo_para = "第二部分  多选题(8题)"
        q_doc.add_paragraph(duo_para)
        a_doc.add_paragraph(duo_para)  # 答案

        add_sheet_part(duo_sheet, q_doc, a_doc, 8, is_mess_up=is_mess_up)

        # ------------------------------*************************------------------------------
        tian_para = "第三部分  填空题(8题)"
        q_doc.add_paragraph(tian_para)
        a_doc.add_paragraph(tian_para)  # 答案

        add_sheet_part(tian_sheet, q_doc, a_doc, 8, is_choice=False)
        # ------------------------------*************************------------------------------
        pan_para = "第四部分  判断题(8题)"
        q_doc.add_paragraph(pan_para)
        a_doc.add_paragraph(pan_para)  # 答案

        add_sheet_part(pan_sheet, q_doc, a_doc, 8, is_choice=False)
        # ------------------------------*************************------------------------------
        jian_para = "第五部分  简答题(2题)"
        q_doc.add_paragraph(jian_para)
        a_doc.add_paragraph(jian_para)  # 答案

        add_sheet_part(jian_sheet, q_doc, a_doc, 2, is_choice=False)
        # ------------------------------*************************------------------------------

        # 保存文件
        q_savepath = os.path.join(output_path, f"测试{count}.docx")
        a_savepath = os.path.join(output_path, f"测试{count}答案.docx")
        q_doc.save(q_savepath)
        a_doc.save(a_savepath)

    wb.close()


if __name__ == "__main__":
    file = "./题库（甲级）.xlsx"
    create(file, 8, False)
